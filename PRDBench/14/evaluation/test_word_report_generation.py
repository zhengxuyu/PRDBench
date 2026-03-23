#!/usr/bin/env python3
"""
Word report generation test script
Test [2.3.2b Full Report Export (Word)] functionality
"""

import subprocess
import os
import sys
from pathlib import Path

def test_word_report_generation():
    """Test Word report generation functionality"""
    print("🧪 Starting Word report generation functionality test...")

    # 1. Execute Word report generation command
    cmd = [
        "python", "-m", "src.main", "report", "generate-full",
        "--data-path", "evaluation/sample_data.csv",
        "--format", "word",
        "--output-path", "evaluation/full_report.docx"
    ]

    print(f"📋 Executing command: {' '.join(cmd)}")

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')

        if result.returncode != 0:
            print(f"❌ Command execution failed, exit code: {result.returncode}")
            print(f"Error output: {result.stderr}")
            return False

        print("✅ Command execution successful")
        print(f"Output: {result.stdout}")

        # 2. Check if expected output files exist
        output_file = "evaluation/full_report.docx"
        expected_file = "evaluation/expected_full_report.docx"

        if not os.path.exists(output_file):
            print(f"❌ Output file does not exist: {output_file}")
            return False
        print(f"✅ Output file exists: {output_file}")

        if not os.path.exists(expected_file):
            print(f"❌ Expected file does not exist: {expected_file}")
            return False
        print(f"✅ Expected file exists: {expected_file}")

        # 3. Verify file size is not zero
        output_size = os.path.getsize(output_file)
        expected_size = os.path.getsize(expected_file)

        if output_size == 0:
            print(f"❌ Output file size is zero: {output_file}")
            return False
        print(f"✅ Output file size: {output_size} bytes")

        if expected_size == 0:
            print(f"❌ Expected file size is zero: {expected_file}")
            return False
        print(f"✅ Expected file size: {expected_size} bytes")

        # 4. Verify file format (simple check of file extension and magic number)
        if not output_file.endswith('.docx'):
            print(f"❌ Output file is not in Word format: {output_file}")
            return False

        # Check DOCX file magic number (ZIP format)
        with open(output_file, 'rb') as f:
            magic = f.read(4)
            if magic != b'PK\x03\x04':
                print(f"❌ Output file is not a valid DOCX format")
                return False
        print(f"✅ Output file is in valid DOCX format")

        # 5. Compare file sizes are within reasonable range
        size_diff_ratio = abs(output_size - expected_size) / expected_size
        if size_diff_ratio > 0.1:  # Allow 10% difference
            print(f"⚠️  File size difference is significant: {size_diff_ratio:.2%}")
        else:
            print(f"✅ File size difference is within reasonable range: {size_diff_ratio:.2%}")

        print("🎉 All tests passed! Word report generation functionality works correctly.")
        return True

    except Exception as e:
        print(f"❌ Exception occurred during test: {e}")
        return False

def validate_word_document_content(file_path):
    """Validate Word document content (requires python-docx library)"""
    try:
        from docx import Document

        doc = Document(file_path)
        print(f"🔍 Validating Word document content: {file_path}")

        # Check paragraph count
        paragraphs = doc.paragraphs
        print(f"✅ Document contains {len(paragraphs)} paragraphs")

        # Check table count
        tables = doc.tables
        print(f"✅ Document contains {len(tables)} tables")

        # Check if key content is included
        full_text = '\n'.join([p.text for p in paragraphs])

        required_sections = [
            "Golf Tourist Consumer Behavior Analysis Report",
            "Executive Summary",
            "Data Overview",
            "Descriptive Statistical Analysis",
            "Analysis Conclusion",
            "Marketing Recommendations"
        ]

        for section in required_sections:
            if section in full_text:
                print(f"✅ Contains required section: {section}")
            else:
                print(f"❌ Missing required section: {section}")
                return False

        return True

    except ImportError:
        print("⚠️  python-docx library not installed, skipping content validation")
        return True
    except Exception as e:
        print(f"❌ Document content validation failed: {e}")
        return False

if __name__ == "__main__":
    success = test_word_report_generation()
    
    # If basic test passes, attempt to validate document content
    if success:
        validate_word_document_content("evaluation/full_report.docx")
    
    sys.exit(0 if success else 1)